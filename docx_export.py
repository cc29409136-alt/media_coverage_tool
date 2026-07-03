"""Word（.docx）匯出模組。

從搜尋結果（original_results / syndication_results，與 matcher.format_report 使用的
資料結構完全相同）產生一份排版過的 Word 文件：標題（關鍵字＋日期區間）、每個媒體一個
子標題、底下是編號清單，每筆標題為粗體文字、網址為可點擊的超連結。

python-docx 目前的版本沒有內建 add_hyperlink()，這裡採用官方 FAQ / 社群常見的
XML workaround 手動組出 <w:hyperlink> 元素。
"""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from matcher import find_origin_match


def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    """在段落中插入一個可點擊的超連結 run（python-docx 標準 workaround）。"""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)

    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _add_media_section(doc, section_title, articles, counter, origin_lookup=None):
    """加入一個媒體（或轉載平台）的子標題＋編號清單，回傳更新後的流水號。"""
    doc.add_heading(section_title, level=2)
    for art in articles:
        p = doc.add_paragraph(style="List Number")
        label = ""
        if origin_lookup is not None:
            origin = find_origin_match(art["title"], origin_lookup)
            if origin:
                label = f"（轉載自：{origin}）"
        run = p.add_run(f"{art['title']}{label}")
        run.bold = True
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Pt(18)
        add_hyperlink(p2, art["url"], art["url"])
        counter += 1
    return counter


def build_docx_report(keyword, start_date, end_date, original_results, syndication_results, original_articles_for_match):
    """產生媒體露出整理的 Word 文件，回傳可直接餵給 st.download_button 的 bytes。"""
    doc = Document()

    date_str = f"{start_date.strftime('%Y/%m/%d')}－{end_date.strftime('%Y/%m/%d')}"
    title = doc.add_heading(f"{keyword}｜{date_str} 媒體露出整理", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    total_found = sum(len(v) for v in original_results.values()) + sum(len(v) for v in syndication_results.values())
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"共找到 {total_found} 篇報導　｜　產出時間：{__import__('datetime').datetime.now().strftime('%Y/%m/%d %H:%M')}")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    counter = 1
    has_original = any(articles for articles in original_results.values())
    if has_original:
        doc.add_heading("原生媒體", level=1)
        for site_name, articles in original_results.items():
            if not articles:
                continue
            counter = _add_media_section(doc, site_name, articles, counter)

    has_syndication = any(articles for articles in syndication_results.values())
    if has_syndication:
        doc.add_heading("轉載平台", level=1)
        for platform_name, articles in syndication_results.items():
            if not articles:
                continue
            counter = _add_media_section(
                doc, f"{platform_name} 轉載", articles, counter, origin_lookup=original_articles_for_match
            )

    if not has_original and not has_syndication:
        doc.add_paragraph("（本次搜尋沒有找到符合關鍵字的報導）")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
